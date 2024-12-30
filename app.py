from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.common.by import By
from datetime import datetime as dt
from docx import Document
from docx.shared import Inches
import os

def start_driver():
    chrome_options = Options()
    arguments = ['--lang=pt-BR', '--start-maximized', '--incognito']

    for argument in arguments:
        chrome_options.add_argument(argument)

    chrome_options.add_experimental_option('prefs', {
        'download.default_directory': 'C:\\Users\\secretario',
        'download.directory_upgrade': True,
        'download.prompt_for_download': False,
        'profile.default_content_setting_values.notifications': 2,
        'profile.default_content_setting_values.automatic_downloads': 1,
    })

    driver = webdriver.Chrome(options=chrome_options)

    return driver

def main():
    driver = start_driver()
    try:
        #Entrando no site
        site= 'https://economia.uol.com.br/cotacoes/cambio/'
        driver.get(site)
        
        graphic = driver.execute_script(f'''
            return document.evaluate(
                '//div[@class= "col-sm-24 col-xs-8 chart-content "]', 
                document,
                null, 
                XPathResult.FIRST_ORDERED_NODE_TYPE, 
                null
                ).singleNodeValue;
            ''')
        driver.execute_script('arguments[0].scrollIntoView({behavior: "instant", block: "center"});', graphic)
        graphic.screenshot('graphic.png')
        
        price= driver.find_element(By.XPATH, '//span[@class= "chart-info-val ng-binding"]').text
        date= dt.now().strftime('%d/%m/%Y')
        img= os.getcwd() + '/graphic.png'

        #Criando documento WORD
        document= Document()

        # Adiciona um título
        document.add_heading('Cotação Atual do Dólar', level=1)

        # Adiciona um parágrafo
        document.add_paragraph(f'''
                                O dólar está no valor de {price}, na data {date}.
                                Valor cotado no site {site}
                                Print da cotação atual:
                            ''')


        # Adiciona uma imagem (certifique-se de que a imagem existe no diretório)
        img = os.getcwd() + '/graphic.png'
        document.add_picture(img, width=Inches(2))

        document.add_paragraph('Cotação feita por - Cauan Neves')

        # Salva o documento
        document.save('doc_cotacao.docx')

        print('Documento Word criado com sucesso!')
        
    finally:
        driver.quit()
    
if __name__ == '__main__':
    main()